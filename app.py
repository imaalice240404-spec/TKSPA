# -*- coding: utf-8 -*-
import os, json, random, time, io, base64, re, tempfile
import streamlit as st
import pandas as pd
import google.generativeai as genai
from supabase import create_client, Client
from docx import Document

# ==========================================
# ⚙️ 1. 系統初始化與環境設定
# ==========================================
st.set_page_config(page_title="研發部專屬 - 專利戰略分析系統", page_icon="⚡", layout="wide")

# 加入 CSS 讓 Streamlit 的按鈕可以支援換行 (變成直的)
st.markdown("""
    <style>
    div[data-testid="stButton"] button {
        white-space: pre-line !important;
        height: 100%;
        min-height: 120px;
    }
    </style>
""", unsafe_allow_html=True)

def get_config(keys):
    for k in keys:
        try: return st.secrets[k]
        except: continue
    return None

S_URL = get_config(["SUPABASE_URL"])
S_KEY = get_config(["SUPABASE_KEY"])
ADMIN_ID = "3676" 

if not S_URL or not S_KEY:
    st.error("🚨 嚴重錯誤：無法從 Secrets 讀取到 Supabase 網址或金鑰！")
    st.stop()

key_pool = []
if get_config(["GOOGLE_API_KEY_1"]): key_pool.append(st.secrets["GOOGLE_API_KEY_1"])
if get_config(["GOOGLE_API_KEY_2"]): key_pool.append(st.secrets["GOOGLE_API_KEY_2"])
if not key_pool and get_config(["GOOGLE_API_KEY"]): key_pool.append(st.secrets["GOOGLE_API_KEY"])

if not key_pool:
    st.error("❌ 系統偵測到 Google API Key 缺失！")
    st.stop()
    
SELECTED_G_KEY = random.choice(key_pool)
genai.configure(api_key=SELECTED_G_KEY)
model = genai.GenerativeModel('gemini-2.5-flash', generation_config=genai.types.GenerationConfig(temperature=0.1, top_p=0.8))

@st.cache_resource
def init_supabase() -> Client:
    return create_client(S_URL, S_KEY)
supabase = init_supabase()

# ==========================================
# 🔐 2. 員工職號登入機制
# ==========================================
if 'current_user' not in st.session_state: 
    st.session_state.current_user = None

if not st.session_state.current_user:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("<br><br><h1 style='text-align: center; color: #1e3a8a;'>⚡ 研發部專利分析系統</h1>", unsafe_allow_html=True)
        with st.container(border=True):
            st.markdown("### 🔐 內部人員登入")
            job_id_input = st.text_input("請輸入您的員工編號：", placeholder="例如：3676")
            if st.button("登入系統", use_container_width=True, type="primary"):
                if job_id_input.strip():
                    st.session_state.current_user = job_id_input.strip()
                    st.rerun()
                else: 
                    st.error("請輸入有效的員工編號！")
    st.stop()

IS_ADMIN = (st.session_state.current_user == ADMIN_ID)

# ==========================================
# 🧠 3. 高階專利 Prompt 庫 
# ==========================================
DETAILED_11_RULES = """
【一、 🚦 FTO 風險判定】
(🔴 紅燈：具威脅 / 🟡 黃燈：需注意 / 🟢 綠燈：已失效。判定原則：若狀態為「公告/核准」或「公開」，絕對不可判定為綠燈！若為「消滅/無效/撤回」才可判為🟢)

【二、📸 技術核心快照】
1. 發明目的：(說明解決傳統弊病) 
2. 核心技術：(說明具體零件結構設計、材料配方或製程步驟) 
3. 宣稱功效：(說明提升了什麼物理效果，如耐壓、高頻特性、降低 ESR 等)

【三、🏢 研發部門精準派發】
[填入建議部門]。 (分發理由)

【四、🛑 先前技術與妥協分析 (防禦地雷)】
本案欲解決之舊設計缺點：(習用技術缺點)
空間配置/製程限制（破口分析）：(列出獨立項限縮最嚴格之特徵或配方參數)

【五、🧩 獨立項全要件拆解 (Claim Chart)】
最廣獨立項（請求項1）拆解：(以 1. 2. 3. 逐行條列拆解)
破口：(精準點出最容易被迴避的限制條件或製程參數)

【六、🪤 附屬項隱藏地雷探測】
(條列出具備具體結構形狀、位置、材料成分比例或參數限制的附屬項)

【七、👁️ 侵權可偵測性評估】
(極易偵測 / 需破壞性拆解如切片(Cross-section)或化學分析(EDX/SEM)，並給出理由)

【八、🕵️‍♂️ 實證功效檢驗 (打假雷達)】
(是否有實體測試數據、可靠度驗證曲線，或僅為定性描述)

【九、🛡️ 高階迴避設計建議 (防範均等論)】
(提出基於破口的具體修改結構方向、替代材料或改變製程工序)

【十、🧬 技術演進與機構整併雷達】
(分析屬於結構整併、材料疊代或製程架構重組)

【十一、 🏷️ IPC 分類號分析】
(列出本案的 IPC 分類號，並簡述其代表的技術領域與分類意義)
"""

PROMPT_M1_BATCH = """
你是一位具備材料科學、化學工程與電子電機碩士學歷，或是在被動元件廠具備實務研發經驗的資深研發主管兼專利工程師。
【⚠️ 跨國語言處理指示】：若輸入的專利標題、摘要或請求項為英文、日文或其他外語，請直接在腦中進行精準的專業翻譯，並【一律使用台灣繁體中文】輸出最終結果！

請嚴格輸出 JSON 格式：
{
  "五大類": "【最高嚴格限制】絕對只能從這 6 個詞彙中挑選：[NTC, PTC, Sensor, Varistor, LCP, 其他]。禁止發明新詞。可多選，用半形逗號分隔。",
  "次系統": "自訂 5-8 字的具體系統名",
  "特殊機構": "15字內精準描述其物理、化學改變或關鍵製程",
  "達成功效": "20字內描述解決的痛點 (例如高頻、高溫或微型化)",
  "核心解法": "用 RD 聽得懂的白話文，極度詳細且精確地描述材料配方、結構疊層、比例參數或製程步驟細節（請務必保留所有具體的數值、材料名稱與特徵條件，字數約 100 到 150 字，提供最詳盡的說明）。"
}
"""

PROMPT_M3_SINGLE = f"""
你是一位具備材料科學、化學工程與電子電機碩士學歷，或是在被動元件廠具備實務經驗的資深研發主管兼專利代理人。
【⚠️ 跨國語言處理指示】：若閱讀的 PDF 說明書為外文（如英文、日文等），請直接將其視為繁體中文進行理解，並【一律使用台灣繁體中文】撰寫下方所有 JSON 內容與報告！
【⚠️ 排版警告】：請絕對不要使用半形波浪號「~」來表示數值範圍（以免導致網頁產生刪除線），請改用「至」或全形「-」。

【🔴 輸出格式要求：純 JSON 格式】
{{
  "rd_card": {{
    "title": "一句話總結", 
    "problem": "傳統缺點", 
    "solution": "本專利特殊結構或製程材料",
    "risk_check": ["1-1. 獨立項全要件限制A", "1-2. 限制B"],
    "design_avoid_rd": ["針對限制A的迴避方向(材料/製程)", "針對限制B的迴避方向"]
  }},
  "vis_data": {{
    "claims": ["1. 獨立項全文...", "2. 依據請求項1..."],
    "components": [ {{"id": "10", "name": "介電層"}} ],
    "spec_texts": ["【00xx】段落內容全文"],
    "loophole_quote": "從請求項1中『一字不漏』複製最能代表本案特徵的那一段。⚠️不含習知技術，標點符號需一致。"
  }},
  "ip_report": "請以專業繁體中文撰寫下方【IP報告十一點】內容。"
}}
【重要】：components 元件與標號必須 100% 精確，絕對不可配錯對。
【IP報告結構】：
{DETAILED_11_RULES}
"""

# ==========================================
# 🛠️ 4. 後端輔助函數
# ==========================================
def parse_ai_json(text):
    try:
        cln = str(text).replace('```json', '').replace('```', '').strip()
        s, e = cln.find('{'), cln.rfind('}')
        return json.loads(cln[s:e+1]) if s != -1 else {}
    except: return {}

def safe_dict(val):
    if isinstance(val, dict): return val
    if isinstance(val, str):
        try: return json.loads(val)
        except: return {}
    return {}

def safe_str(val): return str(val).strip() if pd.notna(val) else ""

# 🌟 解決 Markdown 刪除線的輔助函數
def escape_md(text):
    return str(text).replace("~", "～")

def clean_assignee(name):
    name = safe_str(name)
    if not name: return "未知"
    return re.split(r'股份有限公司|有限公司|公司|Inc|Ltd|LLC|Corporation', name)[0].split(' ')[0].strip() if name else "未知"

def get_db_table(): return 'patents' 

DB_COL_MAP = {
    'id': 'ID', 'app_num': '申請號', 'cert_num': '證書號', 'pub_date': '公開公告日', 'app_date': '申請日',
    'assignee': '專利權人', 'title': '專利名稱', 'abstract': '摘要', 'claims': '請求項',
    'legal_status': '案件狀態', 'status': '狀態', 'sys_main': '五大類', 'sys_sub': '次系統',
    'mechanism': '特殊機構', 'effect': '達成功效', 'solution': '核心解法',
    'thumbnail_base64': '代表圖', 'ipc': 'IPC', 'rd_card_json': 'RDJSON', 'vis_data_json': 'VISJSON', 'ip_report_text': 'REPORT'
}

def fetch_patents(status_filter=None):
    try:
        query = supabase.table(get_db_table()).select("*").order('created_at', desc=True)
        if status_filter: query = query.eq('status', status_filter)
        df = pd.DataFrame(query.execute().data)
        if df.empty: return pd.DataFrame()
        return df.rename(columns=DB_COL_MAP)
    except: return pd.DataFrame()

def create_word_doc(text):
    doc = Document()
    doc.add_heading('專利戰略深度分析報告', 0)
    for para in text.split('\n'):
        if para.strip(): doc.add_paragraph(para.strip())
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def get_patent_type(row):
    cert = str(row.get('證書號', '')).strip().upper()
    app = str(row.get('申請號', '')).strip().upper()
    status = str(row.get('案件狀態', '')).strip()
    
    ref_num = cert if cert else app
    match = re.match(r'^([A-Z]{2})', ref_num)
    country = match.group(1) if match else "TW" 
    
    ptype = "發明" 
    if "新型" in status or "M" in ref_num or ref_num.endswith("U") or ref_num.endswith("Y"):
        ptype = "新型"
    elif "設計" in status or "外觀" in status or "D" in ref_num or ref_num.endswith("S"):
        ptype = "設計"
        
    return f"[{country}] {ptype}"

# ==========================================
# 📊 5. 導覽列與戰略引擎
# ==========================================
for k in ['rd_card_data','claim_data_t2']:
    if k not in st.session_state: st.session_state[k] = {}
if 'target_single_patent' not in st.session_state: st.session_state.target_single_patent = None
if 'ip_report_content' not in st.session_state: st.session_state.ip_report_content = ""
if 'pdf_bytes_main' not in st.session_state: st.session_state.pdf_bytes_main = None
if 'thumbnail_base64' not in st.session_state: st.session_state.thumbnail_base64 = None

PAGES = ["📥 模組一：探勘匯入", "📊 模組二：研發知識庫", "🕵️ 模組三：單篇深度拆解"]

open_count = 0
if IS_ADMIN:
    try: open_count = supabase.table('support_tickets').select('id', count='exact').eq('status', 'OPEN').execute().count or 0
    except: pass
    admin_page_name = f"👑 專家工單中心 (🔴 {open_count} 待處理)" if open_count > 0 else "👑 專家工單中心"
    PAGES.append(admin_page_name)

if 'radio_nav' not in st.session_state: st.session_state.radio_nav = "📊 模組二：研發知識庫"
if 'pending_tab' in st.session_state:
    st.session_state.radio_nav = st.session_state.pending_tab
    del st.session_state.pending_tab
if IS_ADMIN and st.session_state.radio_nav.startswith("👑 專家"):
    st.session_state.radio_nav = admin_page_name

with st.sidebar:
    st.markdown(f"👤 登入職號：**{st.session_state.current_user}**")
    if st.button("🚪 登出系統", use_container_width=True): 
        st.session_state.clear()
        st.rerun()
    st.markdown("---")

    st.radio("功能導覽：", PAGES, key="radio_nav")

    st.markdown("---")
    if st.button("🗑️ 清理當前畫面暫存", use_container_width=True):
        st.session_state.target_single_patent = None
        st.session_state.pdf_bytes_main = None
        for key in ['rd_card_data', 'claim_data_t2']: st.session_state[key] = {}
        st.session_state.ip_report_content = ""
        st.session_state.thumbnail_base64 = None
        st.rerun()

st.title(f"⚡ 被動元件專利 AI 戰略系統 [全球大聲公]")

# ==========================================
# 👑 管理者工單中心
# ==========================================
if st.session_state.radio_nav.startswith("👑 專家"):
    st.header("👑 管理者專屬：專家支援工單中心")
    try:
        tickets = supabase.table('support_tickets').select("*").order('created_at', desc=True).execute().data
        if not tickets: 
            st.success("🎉 目前沒有待處理的工單，大家都很平安！")
        else:
            df_t = pd.DataFrame(tickets)
            open_t = df_t[df_t['status'] == 'OPEN']
            closed_t = df_t[df_t['status'] == 'CLOSED']
            
            t1, t2 = st.tabs([f"🚨 待處理工單 ({len(open_t)})", f"✅ 已結案 ({len(closed_t)})"])
            with t1:
                for _, t in open_t.iterrows():
                    with st.container(border=True):
                        st.markdown(f"### 🎫 工單號: {t['id']} | 專利號: **{t['patent_id']}**")
                        st.caption(f"👤 申請人: {t['job_id']} | 📅 時間: {t['created_at'][:16]}")
                        st.error(f"**🚨 員工疑慮描述：**\n{t['issue_desc']}")
                        ans = st.text_area("✍️ 專業回覆與指導：", key=f"ans_{t['id']}")
                        if st.button("💾 送出回覆並結案", key=f"cls_{t['id']}", type="primary"):
                            supabase.table('support_tickets').update({'admin_reply': ans, 'status': 'CLOSED'}).eq('id', t['id']).execute()
                            st.rerun()
            with t2:
                for _, t in closed_t.iterrows():
                    with st.expander(f"✅ [已結案] 專利號: {t['patent_id']} (申請人: {t['job_id']})"):
                        st.write(f"**問題：** {t['issue_desc']}")
                        st.success(f"**回覆：** {t.get('admin_reply', '無')}")
    except Exception as e: 
        st.error(f"讀取工單失敗: {e}")

# ==========================================
# 📥 模組一：探勘匯入
# ==========================================
elif st.session_state.radio_nav == "📥 模組一：探勘匯入":
    if not IS_ADMIN:
        st.error("⛔ 權限不足：僅限管理員 (3676) 具備資料匯入與探勘權限。")
    else:
        st.header(f"1. 資料匯入與狀態更新 (寫入: `{get_db_table()}`)")
        uploaded_excel = st.file_uploader("上傳 TWPAT/Google Patents 匯出的 Excel/CSV", type=["xlsx", "xls", "csv"])

        if uploaded_excel:
            if st.button("🔄 執行資料比對與匯入", type="primary"):
                df = pd.read_csv(uploaded_excel) if uploaded_excel.name.endswith('.csv') else pd.read_excel(uploaded_excel)
                col_map = {
                    'title': next((c for c in df.columns if '名稱' in c or '標題' in c or 'title' in c.lower()), None),
                    'abs': next((c for c in df.columns if '摘要' in c or 'abstract' in c.lower()), None),
                    'claim': next((c for c in df.columns if '範圍' in c or '請求' in c or 'claim' in c.lower()), None),
                    'app_num': next((c for c in df.columns if '申請號' in c or 'application' in c.lower()), None),
                    'cert_num': next((c for c in df.columns if '證書' in c or '公告' in c or '公開' in c or 'patent' in c.lower() or 'id' in c.lower()), None),
                    'app_date': next((c for c in df.columns if '申請日' in c or 'filed' in c.lower()), None),
                    'pub_date': next((c for c in df.columns if ('公開日' in c or '公告日' in c or 'pub' in c.lower())), None),
                    'assignee': next((c for c in df.columns if '權人' in c or '申請人' in c or 'assignee' in c.lower()), None),
                    'status': next((c for c in df.columns if '狀態' in c or 'status' in c.lower()), None),
                    'ipc': next((c for c in df.columns if 'IPC' in c.upper()), None)
                }
                
                existing_data = supabase.table(get_db_table()).select("id, app_num, legal_status").execute()
                existing_dict = {d['app_num']: d for d in existing_data.data if d['app_num']}

                new_rows_to_insert = []
                update_count, skip_records = 0, 0
                pb = st.progress(0)
                
                for i, row in df.iterrows():
                    app_val = safe_str(row[col_map['app_num']]) if col_map['app_num'] else ""
                    cert_val = safe_str(row[col_map['cert_num']]) if col_map['cert_num'] else ""
                    new_status = safe_str(row[col_map['status']]) if col_map['status'] else "未知"
                    if not app_val and not cert_val: continue 
                    check_val = app_val if app_val else cert_val

                    if check_val in existing_dict:
                        old = existing_dict[check_val]
                        if old['legal_status'] != new_status:
                            upd = {'legal_status': new_status}
                            if "公告" in new_status or "核准" in new_status:
                                upd.update({'cert_num': cert_val, 'rd_card_json': None, 'vis_data_json': None, 'ip_report_text': None})
                            supabase.table(get_db_table()).update(upd).eq('id', old['id']).execute()
                            update_count += 1
                        else: skip_records += 1 
                    else:
                        new_rows_to_insert.append({
                            'app_num': app_val, 'cert_num': cert_val,
                            'app_date': safe_str(row[col_map['app_date']]) if col_map['app_date'] else "未知",
                            'pub_date': safe_str(row[col_map['pub_date']]) if col_map['pub_date'] else "未知",
                            'assignee': clean_assignee(safe_str(row[col_map['assignee']])),
                            'title': safe_str(row[col_map['title']]) if col_map['title'] else "無名稱",
                            'abstract': safe_str(row[col_map['abs']]).replace('\n', '')[:500] if col_map['abs'] else "無摘要",
                            'claims': safe_str(row[col_map['claim']]).replace('\n', '')[:500] if col_map['claim'] else "無請求項",
                            'legal_status': new_status, 'status': 'PENDING',
                            'ipc': safe_str(row[col_map['ipc']]) if col_map['ipc'] else "未知"
                        })
                    if i % 10 == 0: pb.progress(min(1.0, (i + 1) / len(df)))
                
                pb.progress(1.0)
                if new_rows_to_insert:
                    for i in range(0, len(new_rows_to_insert), 500):
                        supabase.table(get_db_table()).insert(new_rows_to_insert[i:i+500]).execute()
                st.success(f"✅ 同步完成！新增: {len(new_rows_to_insert)} | 更新: {update_count} | 跳過: {skip_records}")

        st.markdown("---")
        st.header("2. AI 批次特徵萃取 (支援多國語言)")
        pend_df = fetch_patents('PENDING')
        if not pend_df.empty:
            bs = st.slider("處理筆數", 1, min(50, len(pend_df)), min(5, len(pend_df)))
            if st.button(f"🤖 啟動探勘管線", type="primary"):
                pb2 = st.progress(0)
                for i, (idx, row) in enumerate(pend_df.head(bs).iterrows()):
                    try:
                        res = model.generate_content(f"{PROMPT_M1_BATCH}\n【分析專利】：\n標題:{row['專利名稱']}\n摘要:{row['摘要']}\n請求項:{row['請求項']}").text
                        js = parse_ai_json(res)
                        cats = [c.strip() for c in js.get('五大類', '其他').split(',') if c.strip() in ['NTC', 'PTC', 'Sensor', 'Varistor', 'LCP', '其他']]
                        supabase.table(get_db_table()).update({
                            'sys_main': ', '.join(cats) if cats else '其他', 'sys_sub': js.get('次系統', '未分類'),
                            'mechanism': js.get('特殊機構', ''), 'effect': js.get('達成功效', ''),
                            'solution': js.get('核心解法', ''), 'status': 'COMPLETED'
                        }).eq('id', row['ID']).execute()
                    except Exception as e: 
                        supabase.table(get_db_table()).update({'status': 'FAILED'}).eq('id', row['ID']).execute()
                    pb2.progress((i + 1) / bs)
                    time.sleep(4)
                st.success("✅ 批次解析完成！")
                time.sleep(1)
                st.rerun()

# ==========================================
# 📊 模組二：研發知識庫
# ==========================================
elif st.session_state.radio_nav == "📊 模組二：研發知識庫":
    df = fetch_patents('COMPLETED')
    if df.empty: 
        st.warning("⚠️ 目前資料庫無分析資料，請先至模組一匯入 (或等候 AI 解析)。")
    else:
        df['專利類型'] = df.apply(get_patent_type, axis=1)

        with st.expander("🔍 顯示進階篩選條件"):
            search_q = st.text_input("🔑 關鍵字或號碼搜尋")
            
        fdf = df.copy()
        if search_q:
            qc = re.sub(r'[^a-zA-Z0-9]', '', search_q).upper()
            fdf = fdf[fdf.astype(str).apply(lambda x: x.str.contains(search_q, case=False)).any(axis=1) | 
                      fdf['證書號'].astype(str).str.replace(r'[^a-zA-Z0-9]', '', regex=True).str.upper().str.contains(qc)]

        st.markdown("### 📂 選擇戰略技術區域")
        tech_categories = ["NTC", "PTC", "Sensor", "Varistor", "LCP", "其他"]
        tabs = st.tabs(tech_categories)
        
        for idx, cat in enumerate(tech_categories):
            with tabs[idx]:
                cat_df = fdf[fdf['五大類'].astype(str).str.contains(cat, na=False)]
                if cat_df.empty:
                    st.info(f"💡 目前尚未有 {cat} 相關的專利紀錄。")
                else:
                    for _, p in cat_df.iterrows():
                        did = p['證書號'] if p['證書號'] else p['申請號']
                        # 🌟 防呆：為同一個專利在不同 Tab 中建立唯一 Key，解決 Duplicate Key 錯誤
                        u_key = f"{did}_{idx}"
                        
                        with st.container(border=True):
                            # 🌟 [UI 翻新] 放大圖示比例，按鈕縮小
                            col_img, col_mid, col_btn = st.columns([4, 6.5, 1.5])
                            
                            with col_img:
                                if p.get('代表圖') and len(str(p.get('代表圖'))) > 100: 
                                    st.image(f"data:image/jpeg;base64,{p['代表圖']}", use_container_width=True)
                                else: 
                                    st.markdown("<div style='border:1px dashed #ccc; height:250px; display:flex; align-items:center; justify-content:center; color:#999; background:#fafafa; border-radius:8px;'>🖼️ 無代表圖</div>", unsafe_allow_html=True)
                            
                            with col_mid:
                                st.markdown(f"#### [{did}] {escape_md(p.get('專利名稱', '未知名稱'))}")
                                st.caption(f"🏢 {escape_md(p.get('專利權人', '未知'))} ｜ 📅 日期: {p.get('公開公告日', '未知')} ｜ 🏷️ **{p.get('專利類型', '未知')}**")
                                
                                tags_html = f"""
                                <div style="display:flex; flex-wrap:wrap; gap:10px; margin: 12px 0;">
                                    <div style="background:#e3f2fd; color:#0d47a1; padding:6px 12px; border-radius:6px; font-size:14px;">📁 <b>分類：</b>{escape_md(p.get('五大類', ''))} ➡️ {escape_md(p.get('次系統', ''))}</div>
                                    <div style="background:#fff8e1; color:#f57f17; padding:6px 12px; border-radius:6px; font-size:14px;">⚙️ <b>機構/製程：</b>{escape_md(p.get('特殊機構', '無資料'))}</div>
                                    <div style="background:#fce4ec; color:#c2185b; padding:6px 12px; border-radius:6px; font-size:14px;">🎯 <b>功效：</b>{escape_md(p.get('達成功效', '無資料'))}</div>
                                </div>
                                """
                                st.markdown(tags_html, unsafe_allow_html=True)
                                
                                if p.get('核心解法'): 
                                    st.markdown(f"<div style='color:#444; font-size:15px; line-height:1.6; margin-top:8px;'>💡 **解法細節：** {escape_md(p['核心解法'])}</div>", unsafe_allow_html=True)
                                
                                # 🌟 [新功能] 管理員限定的重新分類工具
                                if IS_ADMIN:
                                    with st.expander("⚙️ 管理員：校正專利分類"):
                                        c_upd1, c_upd2, c_upd3 = st.columns([2, 2, 1])
                                        new_main = c_upd1.multiselect("五大類", tech_categories, default=[c.strip() for c in str(p.get('五大類','')).split(',') if c.strip() in tech_categories], key=f"sel_m_{u_key}")
                                        new_sub = c_upd2.text_input("次系統", value=str(p.get('次系統','')), key=f"sel_s_{u_key}")
                                        if c_upd3.button("💾 儲存分類", key=f"upd_c_{u_key}", use_container_width=True):
                                            supabase.table(get_db_table()).update({'sys_main': ', '.join(new_main) if new_main else '其他', 'sys_sub': new_sub}).eq('id', p['ID']).execute()
                                            st.toast("✅ 分類已更新！")
                                            time.sleep(0.5)
                                            st.rerun()

                            with col_btn:
                                # 🌟 [UI 翻新] 直式按鈕設計
                                if st.button("進\n入\n拆\n解", key=f"btn_s_{u_key}", use_container_width=True, type="primary"):
                                    st.session_state.target_single_patent = p.to_dict()
                                    st.session_state.pdf_bytes_main = None 
                                    for key in ['rd_card_data', 'claim_data_t2']: st.session_state[key] = {}
                                    st.session_state.ip_report_content, st.session_state.thumbnail_base64 = "", None
                                    st.session_state.pending_tab = "🕵️ 模組三：單篇深度拆解"
                                    st.rerun()

# ==========================================
# 🕵️ 模組三：單篇深度拆解
# ==========================================
elif st.session_state.radio_nav == "🕵️ 模組三：單篇深度拆解":
    t = st.session_state.target_single_patent
    if not t: 
        st.warning("👈 請先從模組二選擇一篇專利進入。")
    else:
        db_id, did = t.get('ID'), (t.get('證書號') or t.get('申請號'))
        st.header(f"🕵️ 深度拆解：[{did}] {escape_md(t.get('專利名稱'))}")
        st.markdown(f"**🏢 權利人：** {escape_md(t.get('專利權人'))} | **📅 公開日：** {t.get('公開公告日')} ｜ 🏷️ 類型：**{t.get('專利類型')}**")
        st.markdown("---")

        if not st.session_state.rd_card_data:
            res = supabase.table(get_db_table()).select("rd_card_json, vis_data_json, ip_report_text, thumbnail_base64").eq('id', db_id).execute().data
            if res and res[0].get('rd_card_json'):
                st.session_state.rd_card_data = safe_dict(res[0].get('rd_card_json'))
                st.session_state.claim_data_t2 = safe_dict(res[0].get('vis_data_json'))
                st.session_state.ip_report_content = res[0].get('ip_report_text')
                st.session_state.thumbnail_base64 = res[0].get('thumbnail_base64')

        if IS_ADMIN:
            with st.expander("🛠️ 管理員專區：上傳代表圖與分析文件", expanded=(not st.session_state.rd_card_data)):
                st.info("此處上傳的圖示將直接展示給所有研發人員檢閱。AI 會自動處理多國語言翻譯。")
                col_up1, col_up2 = st.columns(2)
                
                with col_up1:
                    up_img = st.file_uploader("🖼️ 1. 上傳/更新 專利圖示 (PNG/JPG)", type=["png", "jpg", "jpeg"])
                    if up_img and st.button("💾 儲存並更新圖示", use_container_width=True):
                        b64 = base64.b64encode(up_img.getvalue()).decode()
                        supabase.table(get_db_table()).update({'thumbnail_base64': b64}).eq('id', db_id).execute()
                        st.session_state.thumbnail_base64 = b64
                        st.success("✅ 圖示已更新！")
                        st.rerun()

                with col_up2:
                    up_pdf = st.file_uploader("📂 2. 上傳 PDF 執行 11 大天條拆解", type=["pdf"])
                    if up_pdf and st.button("🚀 啟動 AI 深度解析", type="primary", use_container_width=True):
                        st.session_state.pdf_bytes_main = up_pdf.getvalue()
                        with st.spinner("🧠 正在搜索全文並套用 11 大天條..."):
                            try:
                                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp: 
                                    tmp.write(st.session_state.pdf_bytes_main)
                                    tp = tmp.name
                                gf = genai.upload_file(tp)
                                
                                res = model.generate_content([gf, PROMPT_M3_SINGLE])
                                js = parse_ai_json(res.text) 
                                
                                st.session_state.rd_card_data = js.get("rd_card", {})
                                st.session_state.claim_data_t2 = js.get("vis_data", {})
                                st.session_state.ip_report_content = js.get("ip_report", "")
                                
                                upd_payload = {
                                    'rd_card_json': st.session_state.rd_card_data,
                                    'vis_data_json': st.session_state.claim_data_t2,
                                    'ip_report_text': st.session_state.ip_report_content
                                }
                                supabase.table(get_db_table()).update(upd_payload).eq('id', db_id).execute()
                                
                                try: genai.delete_file(gf.name)
                                except: pass 
                                os.remove(tp)
                                st.rerun()
                            except Exception as e: 
                                st.error(f"分析發生異常：{e}")

        if not st.session_state.rd_card_data:
            if not IS_ADMIN:
                st.warning("⏳ 智權主管尚未建立此專利的深度拆解資料。")
        else:
            t_rd, t_ip = st.tabs(["🧑‍💻 Tab 1: 研發迴避大屏", "⚖️ Tab 2: 智權 11 大天條中心"])
            
            with t_rd:
                rd = st.session_state.rd_card_data
                
                # 🌟 上半部：三大看板並列
                c1, c2, c3 = st.columns(3)
                with c1:
                    with st.container(border=True, height=550):
                        st.markdown(f"#### 🎯 研發戰略看板\n**{escape_md(rd.get('title', ''))}**")
                        st.markdown(f"**🔥 解決痛點：**\n\n{escape_md(rd.get('problem', ''))}\n\n**💡 核心解法：**\n\n{escape_md(rd.get('solution', ''))}")
                with c2:
                    with st.container(border=True, height=550):
                        st.markdown("#### 🛡️ 獨立項全要件檢核")
                        st.caption("全要件原則：符合下方所有特徵，則侵權風險極高。")
                        ck_cnt = 0
                        r_list = rd.get('risk_check', [])
                        for i, r in enumerate(r_list):
                            # 加上 escape_md 避免 Markdown 刪除線
                            if st.checkbox(escape_md(str(r)), key=f"rc_{i}"): ck_cnt += 1
                        st.markdown("<br>", unsafe_allow_html=True)
                        if r_list:
                            if ck_cnt == len(r_list): 
                                st.markdown("<div style='padding:10px; background-color:#ffebee; color:#c62828; border-radius:5px;'><b>⚠️ 高度侵權風險！</b></div>", unsafe_allow_html=True)
                            else: 
                                st.markdown("<div style='padding:10px; background-color:#e8f5e9; color:#2e7d32; border-radius:5px;'><b>🎉 文義迴避成功。</b></div>", unsafe_allow_html=True)
                with c3:
                    with st.container(border=True, height=550):
                        st.markdown("#### 🛡️ 高階迴避建議")
                        for a in rd.get('design_avoid_rd', []): 
                            st.markdown(f"✅ {escape_md(a)}")

                st.markdown("---")
                
                # 🌟 下半部：圖示與請求項 (分離排版)
                st.markdown("### 🖼️ 專利核心圖示與獨立項結構")
                img_col, claim_col = st.columns([1, 1.5])
                
                with img_col:
                    with st.container(border=True):
                        if st.session_state.thumbnail_base64:
                            st.image(f"data:image/jpeg;base64,{st.session_state.thumbnail_base64}", use_container_width=True)
                        else:
                            st.info("尚無圖示，請聯繫管理員上傳。")
                
                with claim_col:
                    with st.container(border=True):
                        claims_list = st.session_state.claim_data_t2.get('claims', [])
                        for c in claims_list:
                            st.markdown(f"<div style='font-size:15px; color:#444; line-height:1.8; margin-bottom:10px;'>{escape_md(c)}</div>", unsafe_allow_html=True)

            with t_ip:
                st.markdown("### ⚖️ 智權法務深度報告")
                c_dl, c_dr = st.columns([3, 1])
                c_dl.markdown("以下為嚴格遵守「智權審查 11 大天條」生成的實務報告：")
                c_dr.download_button("📥 下載 Word 報告", data=create_word_doc(st.session_state.ip_report_content), file_name=f"IP_Report_{did}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                with st.container(height=650, border=True): 
                    st.markdown(escape_md(st.session_state.ip_report_content))
                
            st.markdown("---")
            with st.expander("🚨 AI 解析結果不滿意或有侵權疑慮？呼叫管理者支援"):
                issue = st.text_area("描述您需要協助的部分：")
                if st.button("📨 送出支援工單", type="primary"):
                    if issue.strip():
                        supabase.table("support_tickets").insert({"patent_id": did, "job_id": st.session_state.current_user, "issue_desc": issue.strip(), "status": "OPEN"}).execute()
                        st.success("✅ 工單已送出！智權主管將盡快為您解答。")
                    else: 
                        st.error("請填寫描述。")
